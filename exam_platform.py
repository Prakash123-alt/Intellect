import sqlite3
import os
import json
import re
import time
import logging
from datetime import datetime, timedelta
from collections import Counter

DB_PATH = os.path.join(os.path.dirname(__file__), 'data', 'qa_data.db')
UPLOAD_DIR = os.path.join(os.path.dirname(__file__), 'uploads')
EXPORT_DIR = os.path.join(os.path.dirname(__file__), 'exports')

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(EXPORT_DIR, exist_ok=True)

logger = logging.getLogger('exam_platform')


# ─── Helpers ──────────────────────────────────────────────────────────

def _get_conn():
    conn = sqlite3.connect(DB_PATH, timeout=10)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")
    return conn


def _add_column_if_missing(conn, table, column, coltype):
    """ALTER TABLE ... ADD COLUMN, ignoring the error if it already exists."""
    try:
        conn.execute(f"ALTER TABLE {table} ADD COLUMN {column} {coltype}")
    except sqlite3.OperationalError as e:
        if 'duplicate column' not in str(e).lower():
            raise


QUESTION_TYPES = ('mcq', '1mark', '2mark', '5mark', 'mixed')
MARKS_BY_TYPE = {'mcq': 1, '1mark': 1, '2mark': 2, '5mark': 5}


def _parse_json_response(text):
    """Extract JSON from AI response, handling thinking tags and markdown code blocks."""
    logger.info(f"Raw AI response length: {len(text)}")

    # Step 1: Try to get content OUTSIDE thinking tags first
    stripped = re.sub(r'<think>[\s\S]*?</think>', '', text)
    stripped = re.sub(r'<\|think\|>[\s\S]*?<\|/think\|>', '', stripped)
    stripped = re.sub(r'<think>[\s\S]*$', '', stripped)
    stripped = stripped.strip()

    # Step 2: If nothing outside think tags, search for JSON INSIDE them
    if not stripped or not re.search(r'[\[\{]', stripped):
        logger.info("No content outside think tags, extracting JSON from within thinking")
        # Use the full original text to find JSON
        search_text = text
    else:
        search_text = stripped

    # Step 3: Extract JSON from the text
    json_text = _extract_json(search_text)

    # Step 4: Clean and parse
    json_text = re.sub(r',\s*(\]|\})', r'\1', json_text)
    json_text = re.sub(r'//[^\n]*', '', json_text)

    try:
        return json.loads(json_text)
    except json.JSONDecodeError as e:
        logger.error(f"JSON parse failed: {e}")
        logger.error(f"Text being parsed (first 500 chars): {json_text[:500]}")
        lines = json_text.strip().split('\n')
        cleaned = '\n'.join(line for line in lines if line.strip())
        return json.loads(cleaned)


def _extract_json(text):
    """Find and extract the largest valid JSON array or object from text."""
    # Try code block first
    match = re.search(r'```(?:json)?\s*\n?([\s\S]*?)\n?\s*```', text)
    if match:
        candidate = match.group(1).strip()
        try:
            json.loads(candidate)
            return candidate
        except json.JSONDecodeError:
            pass

    # Use bracket matching to find the largest valid JSON structure
    candidates = []

    def find_bracket_matches(open_ch, close_ch):
        """Find all matching bracket ranges in text, respecting strings."""
        ranges = []
        for m in re.finditer(re.escape(open_ch), text):
            start = m.start()
            depth = 0
            in_string = False
            escape = False
            end = None
            for i in range(start, len(text)):
                ch = text[i]
                if escape:
                    escape = False
                    continue
                if ch == '\\' and in_string:
                    escape = True
                    continue
                if ch == '"' and not escape:
                    in_string = not in_string
                    continue
                if in_string:
                    continue
                if ch == open_ch:
                    depth += 1
                elif ch == close_ch:
                    depth -= 1
                    if depth == 0:
                        end = i + 1
                        break
            if end:
                ranges.append((start, end))
        return ranges

    # Collect candidates from both arrays and objects, ordered by size (largest first)
    for start, end in find_bracket_matches('[', ']'):
        candidates.append((start, end))
    for start, end in find_bracket_matches('{', '}'):
        candidates.append((start, end))

    # Sort by length descending - largest is most likely the top-level response
    candidates.sort(key=lambda x: x[1] - x[0], reverse=True)

    for start, end in candidates:
        candidate = text[start:end]
        # Remove trailing commas before final brackets
        cleaned = re.sub(r',\s*\]', ']', candidate)
        cleaned = re.sub(r',\s*\}', '}', cleaned)
        try:
            json.loads(cleaned)
            return cleaned
        except json.JSONDecodeError:
            continue

    # Fallback: regex approach
    match = re.search(r'(\[[\s\S]*\])', text)
    if not match:
        match = re.search(r'(\{[\s\S]*\})', text)
    if match:
        return match.group(1).strip()

    raise ValueError("No JSON structure found in AI response")


def _normalize_questions(data):
    """Normalize quiz questions into a list of dicts with expected keys."""
    # If it's a dict with a list inside, extract the list
    if isinstance(data, dict):
        logger.info(f"Quiz response is a dict with keys: {list(data.keys())}")
        # Try known keys first
        for key in ('questions', 'quiz', 'data', 'items', 'mcq', 'quiz_questions'):
            if key in data and isinstance(data[key], list):
                data = data[key]
                break
        if isinstance(data, dict):
            # Single question wrapped in a dict
            if 'question' in data:
                data = [data]
            else:
                # Try ANY list value in the dict
                for v in data.values():
                    if isinstance(v, list) and len(v) > 0:
                        data = v
                        break

    if not isinstance(data, list):
        raise ValueError(f"Expected a list of questions, got {type(data).__name__}: keys={list(data.keys()) if isinstance(data, dict) else 'N/A'}")

    questions = []
    for item in data:
        if isinstance(item, str):
            continue  # Skip stray strings
        if not isinstance(item, dict):
            continue
        # Normalize options format
        opts = item.get('options', {})
        if isinstance(opts, list) and len(opts) >= 4:
            opts = {'A': opts[0], 'B': opts[1], 'C': opts[2], 'D': opts[3]}
            item['options'] = opts
        elif isinstance(opts, dict):
            # Handle lowercase keys
            normalized = {}
            for k, v in opts.items():
                normalized[k.upper().strip()] = v
            item['options'] = normalized
        # Normalize correct answer
        if 'correct' in item:
            item['correct'] = str(item['correct']).strip().upper()[:1]
        elif 'answer' in item:
            item['correct'] = str(item['answer']).strip().upper()[:1]
        elif 'correct_answer' in item:
            item['correct'] = str(item['correct_answer']).strip().upper()[:1]

        # Reject placeholder/empty content
        qtext = str(item.get('question', '')).strip()
        if not qtext or qtext == '...' or len(qtext) < 5:
            continue
        opts = item.get('options', {})
        if not isinstance(opts, dict) or any(str(v).strip() in ('', '...') for v in opts.values()):
            continue

        if 'question' in item and 'correct' in item:
            questions.append(item)

    if not questions:
        raise ValueError("No valid questions found in AI response")

    return questions


def _normalize_flashcards(data):
    """Normalize flashcards into a list of dicts with front/back keys."""
    if isinstance(data, dict):
        for key in ('flashcards', 'cards', 'data', 'items'):
            if key in data and isinstance(data[key], list):
                data = data[key]
                break

    if not isinstance(data, list):
        raise ValueError(f"Expected a list of flashcards, got {type(data).__name__}")

    cards = []
    for item in data:
        if not isinstance(item, dict):
            continue
        # Handle alternate key names
        front = item.get('front') or item.get('question') or item.get('term', '')
        back = item.get('back') or item.get('answer') or item.get('definition', '')
        # Reject placeholder/empty content
        front = str(front).strip()
        back = str(back).strip()
        if not front or not back or front == '...' or back == '...' or len(front) < 3 or len(back) < 3:
            continue
        cards.append({
            'front': front,
            'back': back,
            'difficulty': item.get('difficulty', 'medium')
        })

    if not cards:
        raise ValueError("No valid flashcards found in AI response")

    return cards


def _ai_call(client, model, system_prompt, user_prompt, retries=2):
    """Make an AI call with retry on failure."""
    last_error = None
    for attempt in range(retries + 1):
        try:
            extra = "\n\nIMPORTANT: Return ONLY valid JSON. No thinking, no explanation, no markdown code blocks." if attempt > 0 else ""
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt + extra}
                ],
                temperature=0.7
            )
            raw = response.choices[0].message.content
            logger.info(f"AI call attempt {attempt + 1} succeeded, response length: {len(raw)}")
            return raw
        except Exception as e:
            last_error = e
            logger.error(f"AI call attempt {attempt + 1} failed: {e}")
            if attempt < retries:
                time.sleep(1)
    raise last_error


# ─── Database Setup ───────────────────────────────────────────────────

def init_exam_db():
    conn = sqlite3.connect(DB_PATH, timeout=10)
    c = conn.cursor()

    c.execute('''CREATE TABLE IF NOT EXISTS subjects (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL UNIQUE,
        description TEXT,
        created_at DATETIME DEFAULT CURRENT_TIMESTAMP
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS pdfs (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        subject TEXT,
        filename TEXT NOT NULL,
        original_name TEXT,
        extracted_text TEXT,
        page_count INTEGER DEFAULT 0,
        page_texts TEXT,
        difficulty_analysis TEXT,
        analyzed_at DATETIME,
        uploaded_at DATETIME DEFAULT CURRENT_TIMESTAMP
    )''')
    _add_column_if_missing(conn, 'pdfs', 'page_texts', 'TEXT')
    _add_column_if_missing(conn, 'pdfs', 'difficulty_analysis', 'TEXT')
    _add_column_if_missing(conn, 'pdfs', 'analyzed_at', 'DATETIME')

    c.execute('''CREATE TABLE IF NOT EXISTS flashcards (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        subject TEXT,
        topic TEXT,
        front TEXT NOT NULL,
        back TEXT NOT NULL,
        difficulty TEXT DEFAULT 'medium',
        times_reviewed INTEGER DEFAULT 0,
        confidence INTEGER DEFAULT 0,
        last_reviewed DATETIME,
        next_review DATETIME,
        created_at DATETIME DEFAULT CURRENT_TIMESTAMP
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS quizzes (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        subject TEXT,
        topic TEXT,
        difficulty TEXT DEFAULT 'medium',
        total_questions INTEGER DEFAULT 0,
        question_type TEXT DEFAULT 'mcq',
        source_pdf_id INTEGER,
        page_range TEXT,
        created_at DATETIME DEFAULT CURRENT_TIMESTAMP
    )''')
    _add_column_if_missing(conn, 'quizzes', 'question_type', "TEXT DEFAULT 'mcq'")
    _add_column_if_missing(conn, 'quizzes', 'source_pdf_id', 'INTEGER')
    _add_column_if_missing(conn, 'quizzes', 'page_range', 'TEXT')

    c.execute('''CREATE TABLE IF NOT EXISTS quiz_questions (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        quiz_id INTEGER NOT NULL,
        question TEXT NOT NULL,
        question_type TEXT DEFAULT 'mcq',
        marks INTEGER DEFAULT 1,
        option_a TEXT,
        option_b TEXT,
        option_c TEXT,
        option_d TEXT,
        correct_answer TEXT,
        model_answer TEXT,
        explanation TEXT,
        FOREIGN KEY (quiz_id) REFERENCES quizzes(id)
    )''')
    _add_column_if_missing(conn, 'quiz_questions', 'question_type', "TEXT DEFAULT 'mcq'")
    _add_column_if_missing(conn, 'quiz_questions', 'marks', 'INTEGER DEFAULT 1')
    _add_column_if_missing(conn, 'quiz_questions', 'model_answer', 'TEXT')

    c.execute('''CREATE TABLE IF NOT EXISTS quiz_attempts (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        quiz_id INTEGER NOT NULL,
        score REAL DEFAULT 0,
        total REAL DEFAULT 0,
        percentage REAL DEFAULT 0,
        time_taken REAL DEFAULT 0,
        answers TEXT,
        weak_topics TEXT,
        attempted_at DATETIME DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (quiz_id) REFERENCES quizzes(id)
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS study_plans (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id TEXT DEFAULT 'default',
        subject TEXT,
        exam_date TEXT,
        daily_hours REAL DEFAULT 2,
        hours_per_day REAL DEFAULT 2,
        plan_type TEXT DEFAULT 'daily',
        plan_data TEXT,
        status TEXT DEFAULT 'active',
        created_at DATETIME DEFAULT CURRENT_TIMESTAMP
    )''')
    _add_column_if_missing(conn, 'study_plans', 'user_id', "TEXT DEFAULT 'default'")
    _add_column_if_missing(conn, 'study_plans', 'daily_hours', 'REAL DEFAULT 2')
    _add_column_if_missing(conn, 'study_plans', 'plan_type', "TEXT DEFAULT 'daily'")

    # ─── Smart Study Planner tables ──────────────────────────────────
    c.execute('''CREATE TABLE IF NOT EXISTS study_tasks (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        plan_id INTEGER NOT NULL,
        topic_name TEXT NOT NULL,
        chapter_name TEXT,
        task_type TEXT DEFAULT 'learn',
        difficulty TEXT DEFAULT 'medium',
        estimated_hours REAL DEFAULT 1,
        scheduled_date TEXT,
        status TEXT DEFAULT 'not_started',
        priority_score REAL DEFAULT 0,
        created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
        updated_at DATETIME DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (plan_id) REFERENCES study_plans(id)
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS uploaded_study_materials (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id TEXT DEFAULT 'default',
        plan_id INTEGER,
        subject TEXT,
        file_name TEXT NOT NULL,
        original_name TEXT,
        file_type TEXT,
        pages INTEGER DEFAULT 0,
        extracted_text TEXT,
        detected_topics TEXT DEFAULT '[]',
        upload_date DATETIME DEFAULT CURRENT_TIMESTAMP
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS study_progress (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id TEXT DEFAULT 'default',
        plan_id INTEGER,
        topic_name TEXT NOT NULL,
        subject TEXT,
        completion_percentage REAL DEFAULT 0,
        last_updated DATETIME DEFAULT CURRENT_TIMESTAMP
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS study_sessions (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        subject TEXT,
        topic TEXT,
        duration_minutes INTEGER DEFAULT 0,
        activity_type TEXT,
        started_at DATETIME DEFAULT CURRENT_TIMESTAMP
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS notes (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        subject TEXT,
        title TEXT,
        content TEXT,
        source_type TEXT DEFAULT 'lecture',
        created_at DATETIME DEFAULT CURRENT_TIMESTAMP
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS topic_mastery (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        subject TEXT,
        topic TEXT,
        mastery_level REAL DEFAULT 0,
        quiz_scores TEXT DEFAULT '[]',
        times_studied INTEGER DEFAULT 0,
        is_weak INTEGER DEFAULT 0,
        last_studied DATETIME,
        updated_at DATETIME DEFAULT CURRENT_TIMESTAMP
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS predictions (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        subject TEXT,
        prediction_data TEXT,
        created_at DATETIME DEFAULT CURRENT_TIMESTAMP
    )''')
    _add_column_if_missing(conn, 'predictions', 'analysis_type', "TEXT DEFAULT 'full'")

    # ─── Exam Prediction Assistant tables ───────────────────────────
    c.execute('''CREATE TABLE IF NOT EXISTS exam_documents (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        subject TEXT NOT NULL,
        doc_type TEXT DEFAULT 'paper',
        filename TEXT NOT NULL,
        original_name TEXT,
        year TEXT,
        extracted_text TEXT,
        chunk_count INTEGER DEFAULT 0,
        uploaded_at DATETIME DEFAULT CURRENT_TIMESTAMP
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS topic_predictions (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        prediction_id INTEGER,
        subject TEXT,
        topic TEXT NOT NULL,
        frequency_score REAL DEFAULT 0,
        confidence_score REAL DEFAULT 0,
        importance_level TEXT DEFAULT 'medium',
        created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (prediction_id) REFERENCES predictions(id)
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS repeated_questions (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        prediction_id INTEGER,
        subject TEXT,
        question TEXT NOT NULL,
        years_found TEXT DEFAULT '[]',
        importance TEXT DEFAULT 'medium',
        created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (prediction_id) REFERENCES predictions(id)
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS predicted_papers (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        prediction_id INTEGER,
        subject TEXT,
        paper_data TEXT,
        created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (prediction_id) REFERENCES predictions(id)
    )''')

    conn.commit()
    conn.close()
    logger.info("Exam platform database initialized")


# ─── PDF Processing ──────────────────────────────────────────────────

def extract_pdf_text(filepath):
    """Extract text from a PDF file using PyMuPDF, falling back to PyPDF2.

    Returns (full_text, page_count, page_texts) where page_texts is a list of
    per-page extracted text (used later for page-range/chapter selection).
    """
    # Try PyMuPDF first (much more robust)
    try:
        import pymupdf
        page_texts = []
        with pymupdf.open(filepath) as doc:
            page_count = len(doc)
            for page in doc:
                page_texts.append((page.get_text() or "").strip())
        extracted = "\n\n".join(page_texts).strip()
        if extracted:
            logger.info(f"PyMuPDF extracted {len(extracted)} chars from {page_count} pages")
            return extracted, page_count, page_texts
    except ImportError:
        logger.warning("PyMuPDF not installed, will try PyPDF2")
    except Exception as e:
        logger.warning(f"PyMuPDF extraction failed: {e}")

    # Fall back to PyPDF2
    try:
        import PyPDF2
        page_texts = []
        page_count = 0
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            page_count = len(reader.pages)
            for page in reader.pages:
                page_texts.append((page.extract_text() or "").strip())
        extracted = "\n\n".join(page_texts).strip()
        if extracted:
            logger.info(f"PyPDF2 extracted {len(extracted)} chars from {page_count} pages")
            return extracted, page_count, page_texts
    except Exception as e:
        logger.error(f"PyPDF2 extraction failed: {e}")

    return "", 0, []


def extract_docx_text(filepath):
    """Extract text from a DOCX file using python-docx."""
    try:
        import docx
        doc = docx.Document(filepath)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(' | '.join(cell.text for cell in row.cells))
        text = '\n'.join(parts).strip()
        return text
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return ""


def extract_txt_text(filepath):
    """Read plain text from a TXT file, trying a couple of encodings."""
    for encoding in ('utf-8', 'latin-1'):
        try:
            with open(filepath, 'r', encoding=encoding) as f:
                return f.read().strip()
        except (UnicodeDecodeError, OSError):
            continue
    return ""


def extract_any_text(filepath, original_name):
    """Extract text from a PDF, DOCX, or TXT file based on its extension.
    Returns (text, page_count).
    """
    ext = original_name.rsplit('.', 1)[-1].lower() if '.' in original_name else ''
    if ext == 'pdf':
        text, page_count, _ = extract_pdf_text(filepath)
        return text, page_count
    if ext == 'docx':
        return extract_docx_text(filepath), 0
    if ext == 'txt':
        return extract_txt_text(filepath), 0
    return "", 0


def save_pdf(subject, filename, original_name, text, page_count, page_texts=None):
    conn = _get_conn()
    c = conn.cursor()
    c.execute('''INSERT INTO pdfs (subject, filename, original_name, extracted_text, page_count, page_texts)
        VALUES (?, ?, ?, ?, ?, ?)''',
        (subject, filename, original_name, text, page_count,
         json.dumps(page_texts) if page_texts else None))
    pdf_id = c.lastrowid
    conn.commit()
    conn.close()
    return pdf_id


def get_pdfs(subject=None):
    conn = _get_conn()
    c = conn.cursor()
    if subject:
        c.execute("SELECT * FROM pdfs WHERE subject = ? ORDER BY uploaded_at DESC", (subject,))
    else:
        c.execute("SELECT * FROM pdfs ORDER BY uploaded_at DESC")
    rows = [dict(r) for r in c.fetchall()]
    conn.close()
    return rows


def get_pdf(pdf_id):
    conn = _get_conn()
    c = conn.cursor()
    c.execute("SELECT * FROM pdfs WHERE id = ?", (pdf_id,))
    row = c.fetchone()
    conn.close()
    return dict(row) if row else None


def get_pdf_text_range(pdf, start_page=None, end_page=None):
    """Return the text of a PDF, optionally sliced to a 1-indexed page range."""
    if not pdf:
        return ""
    page_texts = None
    if pdf.get('page_texts'):
        try:
            page_texts = json.loads(pdf['page_texts'])
        except (TypeError, ValueError):
            page_texts = None

    if not start_page and not end_page:
        return pdf.get('extracted_text', '')

    if not page_texts:
        # No per-page text stored (e.g. older upload) - fall back to full text
        return pdf.get('extracted_text', '')

    total_pages = len(page_texts)
    start = max(1, int(start_page)) if start_page else 1
    end = min(total_pages, int(end_page)) if end_page else total_pages
    if start > end:
        start, end = end, start
    selected = page_texts[start - 1:end]
    return "\n\n".join(selected).strip()


# ─── Difficulty Analysis (Simplify Hard Topics) ──────────────────────

def analyze_pdf_difficulty(client, model, subject, text, max_concepts=8):
    """Scan a PDF's text for difficult concepts, technical terms, and complex definitions.
    For each, generate a simplified explanation, a real-world example, and a memory trick."""
    prompt = f"""You are helping a student study "{subject}" from the material below. Identify the concepts,
technical terms, and complex definitions that students usually find difficult to understand.

Material excerpt:
{text[:6000]}

Return a JSON array of up to {max_concepts} objects, each with:
- "concept": the difficult term/concept name (as it appears in the material)
- "category": one of "difficult_concept", "technical_term", "complex_definition"
- "simplified_explanation": a plain-language explanation a beginner can understand (2-3 sentences)
- "real_world_example": a relatable real-world analogy or example
- "memory_trick": a short mnemonic or memory trick to help remember it

Return ONLY the JSON array, no other text."""
    system = ("You are an expert teacher who specializes in simplifying difficult academic material for students. "
              "Respond with ONLY a valid JSON array. Do not include any explanation, commentary, or thinking.")

    for attempt in range(3):
        try:
            raw = _ai_call(client, model, system, prompt)
            result = _parse_json_response(raw)
            if isinstance(result, dict):
                for key in ('concepts', 'items', 'data'):
                    if key in result and isinstance(result[key], list):
                        result = result[key]
                        break
            if not isinstance(result, list):
                raise ValueError("Expected a JSON array of difficult concepts")
            cleaned = []
            for item in result:
                if not isinstance(item, dict):
                    continue
                concept = str(item.get('concept', '')).strip()
                explanation = str(item.get('simplified_explanation', '')).strip()
                if not concept or not explanation:
                    continue
                category = str(item.get('category', 'difficult_concept')).strip().lower()
                if category not in ('difficult_concept', 'technical_term', 'complex_definition'):
                    category = 'difficult_concept'
                cleaned.append({
                    'concept': concept,
                    'category': category,
                    'simplified_explanation': explanation,
                    'real_world_example': str(item.get('real_world_example', '')).strip(),
                    'memory_trick': str(item.get('memory_trick', '')).strip(),
                })
            if cleaned:
                return cleaned
            raise ValueError("No valid difficult concepts found in AI response")
        except (json.JSONDecodeError, ValueError, TypeError) as e:
            logger.warning(f"Difficulty analysis parse attempt {attempt + 1} failed: {e}")
            if attempt == 2:
                raise
            time.sleep(0.5)


def save_pdf_difficulty_analysis(pdf_id, analysis):
    conn = _get_conn()
    conn.execute("UPDATE pdfs SET difficulty_analysis = ?, analyzed_at = ? WHERE id = ?",
                 (json.dumps(analysis), datetime.now().isoformat(), pdf_id))
    conn.commit()
    conn.close()


# ─── Flashcards ───────────────────────────────────────────────────────

def generate_flashcards(client, model, topic, count=10, context=None):
    """Generate flashcards using AI with retry on parse failure."""
    ctx = f"\n\nUse this study material as reference:\n{context[:3000]}" if context else ""
    prompt = f"""Generate exactly {count} flashcards on the topic "{topic}".{ctx}

Return a JSON array where each element has:
- "front": the question or term (concise)
- "back": the answer or definition (clear, 1-3 sentences)
- "difficulty": one of "easy", "medium", "hard"

Return ONLY the JSON array, no other text."""

    system = "You are an expert educator. Generate high-quality flashcards. Respond with ONLY a valid JSON array. Do not include any explanation, commentary, or thinking."

    for attempt in range(3):
        try:
            raw = _ai_call(client, model, system, prompt)
            result = _parse_json_response(raw)
            return _normalize_flashcards(result)
        except (json.JSONDecodeError, ValueError, TypeError) as e:
            logger.warning(f"Flashcard generation parse attempt {attempt + 1} failed: {e}")
            if attempt == 2:
                raise
            time.sleep(0.5)


def save_flashcards(subject, topic, cards):
    conn = _get_conn()
    c = conn.cursor()
    now = datetime.now().isoformat()
    saved = 0
    for card in cards:
        c.execute('''INSERT INTO flashcards (subject, topic, front, back, difficulty, next_review)
            VALUES (?, ?, ?, ?, ?, ?)''',
            (subject, topic, card['front'], card['back'],
             card.get('difficulty', 'medium'), now))
        saved += 1
    conn.commit()
    conn.close()
    return saved


def get_flashcards(subject=None, topic=None):
    conn = _get_conn()
    c = conn.cursor()
    query = "SELECT * FROM flashcards WHERE 1=1"
    params = []
    if subject:
        query += " AND subject = ?"
        params.append(subject)
    if topic:
        query += " AND topic = ?"
        params.append(topic)
    query += " ORDER BY created_at DESC"
    c.execute(query, params)
    rows = [dict(r) for r in c.fetchall()]
    conn.close()
    return rows


def get_due_flashcards(limit=20):
    conn = _get_conn()
    c = conn.cursor()
    now = datetime.now().isoformat()
    c.execute("""SELECT * FROM flashcards 
        WHERE next_review IS NULL OR next_review <= ?
        ORDER BY confidence ASC, next_review ASC LIMIT ?""", (now, limit))
    rows = [dict(r) for r in c.fetchall()]
    conn.close()
    return rows


def update_flashcard_review(card_id, confidence):
    """Update flashcard after review. Confidence: 1-5."""
    intervals = {1: 1, 2: 1, 3: 3, 4: 7, 5: 14}
    days = intervals.get(confidence, 1)
    next_review = (datetime.now() + timedelta(days=days)).isoformat()
    conn = _get_conn()
    c = conn.cursor()
    c.execute("""UPDATE flashcards SET 
        confidence = ?, times_reviewed = times_reviewed + 1,
        last_reviewed = ?, next_review = ?
        WHERE id = ?""", (confidence, datetime.now().isoformat(), next_review, card_id))
    conn.commit()
    conn.close()


def get_flashcard_stats():
    conn = _get_conn()
    c = conn.cursor()
    c.execute("SELECT COUNT(*) FROM flashcards")
    total = c.fetchone()[0]
    now = datetime.now().isoformat()
    c.execute("SELECT COUNT(*) FROM flashcards WHERE next_review IS NULL OR next_review <= ?", (now,))
    due = c.fetchone()[0]
    c.execute("SELECT AVG(confidence) FROM flashcards WHERE confidence > 0")
    avg_conf = c.fetchone()[0] or 0
    c.execute("SELECT COUNT(*) FROM flashcards WHERE confidence >= 4")
    mastered = c.fetchone()[0]
    conn.close()
    return {'total': total, 'due': due, 'avg_confidence': round(avg_conf, 1), 'mastered': mastered}


# ─── Quizzes ──────────────────────────────────────────────────────────

def _normalize_exam_questions(data, question_type):
    """Normalize AI output for any question type (mcq/1mark/2mark/5mark/mixed)
    into a list of dicts with a consistent shape:
    {question, type, marks, options?, correct?, explanation?, model_answer?}
    """
    if isinstance(data, dict):
        for key in ('questions', 'quiz', 'data', 'items'):
            if key in data and isinstance(data[key], list):
                data = data[key]
                break
    if not isinstance(data, list):
        raise ValueError(f"Expected a list of questions, got {type(data).__name__}")

    questions = []
    for item in data:
        if not isinstance(item, dict):
            continue
        qtext = str(item.get('question', '')).strip()
        if not qtext or qtext == '...' or len(qtext) < 5:
            continue

        qtype = str(item.get('type', question_type)).strip().lower() if question_type == 'mixed' else question_type
        if qtype not in MARKS_BY_TYPE:
            qtype = 'mcq'
        marks = MARKS_BY_TYPE[qtype]

        entry = {'question': qtext, 'type': qtype, 'marks': marks}

        if qtype == 'mcq':
            opts = item.get('options', {})
            if isinstance(opts, list) and len(opts) >= 4:
                opts = {'A': opts[0], 'B': opts[1], 'C': opts[2], 'D': opts[3]}
            elif isinstance(opts, dict):
                opts = {k.upper().strip(): v for k, v in opts.items()}
            if not isinstance(opts, dict) or any(str(v).strip() in ('', '...') for v in opts.values()) or len(opts) < 4:
                continue
            correct = item.get('correct') or item.get('answer') or item.get('correct_answer', '')
            correct = str(correct).strip().upper()[:1]
            if not correct:
                continue
            entry['options'] = opts
            entry['correct'] = correct
            entry['explanation'] = str(item.get('explanation', '')).strip()
        else:
            model_answer = str(item.get('model_answer') or item.get('answer', '')).strip()
            if not model_answer or model_answer == '...':
                continue
            entry['model_answer'] = model_answer
            entry['explanation'] = str(item.get('explanation', '')).strip()

        questions.append(entry)

    if not questions:
        raise ValueError("No valid questions found in AI response")

    return questions


_TYPE_PROMPTS = {
    'mcq': """Generate exactly {count} multiple choice questions on "{topic}" at {difficulty} difficulty level.{ctx}

Return a JSON array where each element has:
- "question": the question text
- "options": {{"A": "...", "B": "...", "C": "...", "D": "..."}}
- "correct": the correct option letter (A, B, C, or D)
- "explanation": brief explanation of the correct answer (1-2 sentences)""",

    '1mark': """Generate exactly {count} 1-mark questions on "{topic}" at {difficulty} difficulty level.{ctx}
1-mark questions are very short-answer, definition-based, or key-concept questions that can be answered in one line.

Return a JSON array where each element has:
- "question": the question text
- "model_answer": the ideal concise one-line answer
- "explanation": (optional) a short note on what to look for in a correct answer""",

    '2mark': """Generate exactly {count} 2-mark questions on "{topic}" at {difficulty} difficulty level.{ctx}
2-mark questions require a short explanation demonstrating concept understanding (2-4 sentences).

Return a JSON array where each element has:
- "question": the question text
- "model_answer": the ideal answer (2-4 sentences covering the key points)
- "explanation": (optional) key points the answer should include""",

    '5mark': """Generate exactly {count} 5-mark questions on "{topic}" at {difficulty} difficulty level.{ctx}
5-mark questions are detailed, descriptive, analytical, exam-oriented questions requiring an in-depth answer.

Return a JSON array where each element has:
- "question": the question text
- "model_answer": a detailed model answer covering all key points (can use multiple sentences/paragraphs)
- "explanation": (optional) the key points/structure a full-marks answer should cover""",

    'mixed': """Generate exactly {count} exam questions on "{topic}" at {difficulty} difficulty level, using a balanced mix of question types.{ctx}
Distribute the questions roughly evenly across these types: "mcq" (multiple choice), "1mark" (very short answer), "2mark" (short explanation), "5mark" (detailed/descriptive).

Return a JSON array where each element has:
- "question": the question text
- "type": one of "mcq", "1mark", "2mark", "5mark"
- If type is "mcq": also include "options": {{"A": "...", "B": "...", "C": "...", "D": "..."}} and "correct": the correct option letter, and "explanation"
- If type is "1mark", "2mark", or "5mark": also include "model_answer" (ideal answer, length appropriate to the mark value) and optionally "explanation" """,
}


def generate_exam_quiz(client, model, topic, question_type='mcq', difficulty='medium', count=10, context=None):
    """Generate quiz questions of a given type using AI, with retry on parse failure."""
    if question_type not in QUESTION_TYPES:
        question_type = 'mcq'
    ctx = f"\n\nBase questions on this material:\n{context[:6000]}" if context else ""
    body = _TYPE_PROMPTS[question_type].format(topic=topic, difficulty=difficulty, count=count, ctx=ctx)
    prompt = f"""{body}

Return ONLY the JSON array, no other text."""

    system = ("You are an expert exam question creator. Generate clear, exam-quality questions "
              "appropriate for the requested type and difficulty. Respond with ONLY a valid JSON array. "
              "Do not include any explanation, commentary, or thinking.")

    for attempt in range(3):
        try:
            raw = _ai_call(client, model, system, prompt)
            result = _parse_json_response(raw)
            return _normalize_exam_questions(result, question_type)
        except (json.JSONDecodeError, ValueError, TypeError) as e:
            logger.warning(f"Quiz generation parse attempt {attempt + 1} failed: {e}")
            if attempt == 2:
                raise
            time.sleep(0.5)


def generate_quiz(client, model, topic, difficulty='medium', count=10, context=None):
    """Backward-compatible MCQ-only quiz generator."""
    return generate_exam_quiz(client, model, topic, 'mcq', difficulty, count, context)


def save_quiz(subject, topic, difficulty, questions, question_type='mcq', source_pdf_id=None, page_range=None):
    conn = _get_conn()
    c = conn.cursor()
    c.execute('''INSERT INTO quizzes (subject, topic, difficulty, total_questions, question_type, source_pdf_id, page_range)
        VALUES (?, ?, ?, ?, ?, ?, ?)''',
        (subject, topic, difficulty, len(questions), question_type, source_pdf_id, page_range))
    quiz_id = c.lastrowid
    for q in questions:
        qtype = q.get('type', question_type)
        marks = q.get('marks', MARKS_BY_TYPE.get(qtype, 1))
        opts = q.get('options', {})
        c.execute('''INSERT INTO quiz_questions
            (quiz_id, question, question_type, marks, option_a, option_b, option_c, option_d,
             correct_answer, model_answer, explanation)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''',
            (quiz_id, q['question'], qtype, marks,
             opts.get('A', ''), opts.get('B', ''), opts.get('C', ''), opts.get('D', ''),
             q.get('correct', ''), q.get('model_answer', ''), q.get('explanation', '')))
    conn.commit()
    conn.close()
    return quiz_id


def get_all_quizzes():
    conn = _get_conn()
    c = conn.cursor()
    c.execute("SELECT * FROM quizzes ORDER BY created_at DESC")
    rows = [dict(r) for r in c.fetchall()]
    conn.close()
    return rows


def get_quiz(quiz_id):
    conn = _get_conn()
    c = conn.cursor()
    c.execute("SELECT * FROM quizzes WHERE id = ?", (quiz_id,))
    quiz = c.fetchone()
    if not quiz:
        conn.close()
        return None, []
    c.execute("SELECT * FROM quiz_questions WHERE quiz_id = ?", (quiz_id,))
    questions = [dict(r) for r in c.fetchall()]
    conn.close()
    return dict(quiz), questions


def _grade_subjective_answers(client, model, subjective):
    """Use AI to grade a batch of subjective answers against their model answers.
    subjective: list of dicts {question_id, question, model_answer, marks, user_answer}
    Returns dict {question_id: {'score': float, 'feedback': str}}
    """
    if not subjective:
        return {}
    if not client:
        return {s['question_id']: {'score': 0, 'feedback': 'Not auto-graded. Compare your answer with the model answer above.'}
                for s in subjective}

    items = [{
        'question_id': s['question_id'],
        'question': s['question'],
        'model_answer': s['model_answer'],
        'max_marks': s['marks'],
        'student_answer': s['user_answer'] or '(no answer provided)'
    } for s in subjective]

    prompt = f"""Grade the following student answers against the model answers. Award partial credit where appropriate.

{json.dumps(items, indent=2)}

Return a JSON array where each element has:
- "question_id": matching the input
- "score": marks awarded (a number from 0 to max_marks, can be fractional)
- "feedback": a brief 1-2 sentence explanation of the score

Return ONLY the JSON array, no other text."""

    system = ("You are a strict but fair exam grader. Award marks based on how well the student's answer "
              "covers the key points of the model answer. Respond with ONLY a valid JSON array.")

    try:
        raw = _ai_call(client, model, system, prompt)
        result = _parse_json_response(raw)
        graded = {}
        for item in result:
            qid = item.get('question_id')
            if qid is None:
                continue
            score = item.get('score', 0)
            try:
                score = float(score)
            except (TypeError, ValueError):
                score = 0
            graded[qid] = {'score': score, 'feedback': str(item.get('feedback', '')).strip()}
        return graded
    except Exception as e:
        logger.error(f"Subjective grading failed: {e}")
        return {s['question_id']: {'score': 0, 'feedback': 'Auto-grading failed. Compare your answer with the model answer above.'}
                for s in subjective}


def submit_quiz(quiz_id, answers, time_taken=0, client=None, model=None):
    """Score a quiz attempt. answers = {question_id: selected_letter_or_text}"""
    quiz, questions = get_quiz(quiz_id)
    if not quiz:
        return None

    score = 0.0
    total = 0.0
    wrong_topics = []
    details = []
    subjective = []

    for q in questions:
        qid = q['id']
        qtype = q.get('question_type', 'mcq')
        marks = q.get('marks', 1) or 1
        total += marks
        user_answer = answers.get(str(qid), '')

        if qtype == 'mcq':
            correct = q['correct_answer']
            is_correct = user_answer.strip().upper() == (correct or '').strip().upper()
            obtained = marks if is_correct else 0
            score += obtained
            if not is_correct:
                wrong_topics.append(quiz.get('topic', 'General'))
            details.append({
                'question_id': qid,
                'question_type': qtype,
                'selected': user_answer,
                'correct': correct,
                'is_correct': is_correct,
                'marks': marks,
                'score': obtained
            })
        else:
            subjective.append({
                'question_id': qid,
                'question': q['question'],
                'model_answer': q.get('model_answer', ''),
                'marks': marks,
                'user_answer': user_answer
            })

    graded = _grade_subjective_answers(client, model, subjective)
    for s in subjective:
        qid = s['question_id']
        result = graded.get(qid, {'score': 0, 'feedback': ''})
        obtained = max(0, min(s['marks'], result['score']))
        score += obtained
        is_correct = obtained >= s['marks'] * 0.6
        if not is_correct:
            wrong_topics.append(quiz.get('topic', 'General'))
        details.append({
            'question_id': qid,
            'question_type': next(q['question_type'] for q in questions if q['id'] == qid),
            'selected': s['user_answer'],
            'correct': s['model_answer'],
            'is_correct': is_correct,
            'marks': s['marks'],
            'score': obtained,
            'feedback': result.get('feedback', '')
        })

    percentage = round((score / total * 100), 1) if total > 0 else 0

    conn = _get_conn()
    c = conn.cursor()
    c.execute('''INSERT INTO quiz_attempts 
        (quiz_id, score, total, percentage, time_taken, answers, weak_topics)
        VALUES (?, ?, ?, ?, ?, ?, ?)''',
        (quiz_id, score, total, percentage, time_taken,
         json.dumps(details), json.dumps(wrong_topics)))
    attempt_id = c.lastrowid
    conn.commit()
    conn.close()

    # Update topic mastery
    if quiz.get('topic') and quiz.get('subject'):
        _update_mastery(quiz['subject'], quiz['topic'], percentage)

    return attempt_id


def get_quiz_result(attempt_id):
    conn = _get_conn()
    c = conn.cursor()
    c.execute("SELECT * FROM quiz_attempts WHERE id = ?", (attempt_id,))
    attempt = c.fetchone()
    if not attempt:
        conn.close()
        return None, None, []
    attempt = dict(attempt)
    c.execute("SELECT * FROM quizzes WHERE id = ?", (attempt['quiz_id'],))
    quiz = dict(c.fetchone())
    c.execute("SELECT * FROM quiz_questions WHERE quiz_id = ?", (attempt['quiz_id'],))
    questions = [dict(r) for r in c.fetchall()]
    conn.close()
    attempt['answers'] = json.loads(attempt['answers']) if attempt['answers'] else []
    return attempt, quiz, questions


def get_quiz_history():
    conn = _get_conn()
    c = conn.cursor()
    c.execute("""SELECT qa.*, q.topic, q.subject, q.difficulty 
        FROM quiz_attempts qa JOIN quizzes q ON qa.quiz_id = q.id
        ORDER BY qa.attempted_at DESC LIMIT 20""")
    rows = [dict(r) for r in c.fetchall()]
    conn.close()
    return rows


# ─── Smart Study Planner ─────────────────────────────────────────────

DIFFICULTY_BASE_HOURS = {'easy': 1.0, 'medium': 1.75, 'hard': 2.5}


def _detect_headings(text, max_topics=40):
    """Heuristic chapter/heading/topic detector (no AI) - looks for
    'Chapter/Unit/Module N: ...' lines, numbered headings, and short title-case lines."""
    numbered = re.compile(r'^\s*(chapter|unit|module|lesson)\s+\d+[:\-.]?\s*(.+)$', re.I)
    dotted = re.compile(r'^\s*\d+(\.\d+)*\s+[A-Z][A-Za-z0-9 ,&/\-]{3,70}$')

    candidates = []
    for line in text.split('\n'):
        line = line.strip()
        if not line or len(line) > 90:
            continue
        m = numbered.match(line)
        if m:
            candidates.append(m.group(2).strip() or line)
            continue
        if dotted.match(line):
            candidates.append(line)
            continue
        words = line.split()
        if 2 <= len(words) <= 8 and line[-1] not in '.,;:' and line[0].isupper() and not line.isupper():
            candidates.append(line)

    seen, unique = set(), []
    for c in candidates:
        key = c.lower()
        if key not in seen:
            seen.add(key)
            unique.append(c)
    return unique[:max_topics]


def extract_material_topics(client, model, subject, text):
    """Detect chapters/headings/topics from a study material's text.
    Falls back to an AI call if the heuristic heading detector finds too little."""
    candidates = _detect_headings(text)
    if len(candidates) >= 3:
        return candidates

    prompt = f"""Extract the main chapters/topics/subtopics covered in this study material for "{subject}".

Excerpt:
{text[:5000]}

Return a JSON array of up to 15 short topic/chapter name strings only. Return ONLY the JSON array, no other text."""
    system = ("You are an expert at identifying syllabus topics from study material. "
              "Respond with ONLY a valid JSON array of strings.")
    try:
        result = _ai_json_call(client, model, system, prompt)
        if isinstance(result, list):
            extracted = [str(t).strip() for t in result if str(t).strip()][:15]
            if extracted:
                return extracted
    except Exception as e:
        logger.warning(f"AI topic extraction fallback failed: {e}")
    return candidates


def classify_topics(client, model, subject, topics, context_snippet=None):
    """Classify a list of topic names by difficulty (easy/medium/hard) and importance (0-100)."""
    topics = list(dict.fromkeys(t.strip() for t in topics if t and t.strip()))[:40]
    if not topics:
        return []

    ctx = f"\n\nReference material excerpt:\n{context_snippet[:2000]}" if context_snippet else ""
    prompt = f"""Classify each of the following topics for the subject "{subject}" by difficulty and exam importance.{ctx}

Topics: {json.dumps(topics)}

Difficulty guide:
- "easy": definitions, basic concepts
- "medium": algorithms, processes, standard procedures
- "hard": numerical problems, complex theories, derivations

Return a JSON array where each element has:
- "topic": the exact topic name from the input list
- "difficulty": "easy", "medium", or "hard"
- "importance_score": 0-100 (how central/likely to be tested)

Return ONLY the JSON array, no other text."""
    system = ("You are an expert academic advisor who assesses topic difficulty and exam importance. "
              "Respond with ONLY a valid JSON array.")

    try:
        result = _ai_json_call(client, model, system, prompt)
    except Exception as e:
        logger.warning(f"Topic classification failed, using defaults: {e}")
        result = []

    classified = {}
    if isinstance(result, list):
        for item in result:
            if not isinstance(item, dict):
                continue
            name = str(item.get('topic', '')).strip()
            if not name:
                continue
            difficulty = str(item.get('difficulty', 'medium')).strip().lower()
            if difficulty not in DIFFICULTY_BASE_HOURS:
                difficulty = 'medium'
            try:
                importance = float(item.get('importance_score', 50))
            except (TypeError, ValueError):
                importance = 50
            classified[name.lower()] = {'topic': name, 'difficulty': difficulty, 'importance_score': importance}

    # Ensure every input topic is represented, even if the AI dropped some
    out = []
    for t in topics:
        out.append(classified.get(t.lower(), {'topic': t, 'difficulty': 'medium', 'importance_score': 50}))
    return out


def _prediction_priority_boost(subject):
    """Build a {topic_name_lower: boost_score} map from the latest exam prediction for this subject."""
    preds = get_predictions(subject)
    if not preds:
        return {}
    data = preds[0].get('prediction_data', {})
    boost = {}
    for t in data.get('frequently_asked_topics', []) or []:
        name = str(t.get('topic', '')).strip().lower()
        if name:
            boost[name] = max(boost.get(name, 0), t.get('importance_score', 50) or 0)
    for c in data.get('high_weightage_concepts', []) or []:
        name = str(c.get('concept', '')).strip().lower()
        if name:
            boost[name] = max(boost.get(name, 0), c.get('exam_frequency', 50) or 0)
    return boost


def _match_boost(topic_name, boost_map):
    tl = topic_name.lower()
    best = 0
    for key, val in boost_map.items():
        if key in tl or tl in key:
            best = max(best, val)
    return best


def build_study_schedule(topics, exam_date, daily_hours, plan_type='daily', weak_topic_names=None):
    """Rule-based scheduler: distributes topics across available days, adds spaced-repetition
    revisions and periodic practice tests. Returns (tasks, meta).
    """
    today = datetime.now().date()
    try:
        exam_d = datetime.strptime(exam_date, '%Y-%m-%d').date()
    except ValueError:
        exam_d = today + timedelta(days=14)

    days_remaining = max(1, (exam_d - today).days)
    weak_topic_names = {t.lower() for t in (weak_topic_names or [])}

    # Priority: frequently-asked/high-weightage boost (already merged into 'boost' by the caller),
    # then difficulty, then weak-topic bump
    ranked = []
    for t in topics:
        priority = float(t.get('importance_score', 50)) + float(t.get('boost', 0))
        if t['difficulty'] == 'hard':
            priority += 8
        if t['topic'].lower() in weak_topic_names:
            priority += 15
        ranked.append({**t, 'priority_score': round(priority, 1)})
    ranked.sort(key=lambda x: x['priority_score'], reverse=True)

    # Reserve the final day(s) for consolidated final revision
    final_revision_days = 1 if days_remaining <= 3 else 2
    learn_window = max(1, days_remaining - final_revision_days)
    learn_day_budget = daily_hours * 0.65

    tasks = []
    warnings = []
    day_offset = 1  # 1 = tomorrow relative to "today" (today is left for setup/reading this plan)
    day_load = 0.0

    def add_task(topic_name, chapter_name, task_type, difficulty, hours, offset, priority_score):
        d = today + timedelta(days=max(0, offset))
        if d > exam_d:
            d = exam_d
        tasks.append({
            'topic_name': topic_name, 'chapter_name': chapter_name, 'task_type': task_type,
            'difficulty': difficulty, 'estimated_hours': round(hours, 2),
            'scheduled_date': d.isoformat(), 'priority_score': priority_score
        })

    for t in ranked:
        hours = DIFFICULTY_BASE_HOURS.get(t['difficulty'], 1.75)
        if day_offset > learn_window:
            # Out of learning days - compress onto the last learning day
            warnings.append(f"Not enough days before the exam to fully space out '{t['topic']}'; compressed into the schedule.")
            day_offset = learn_window
            day_load = 0.0
        if day_load + hours > learn_day_budget and day_load > 0:
            day_offset += 1
            day_load = 0.0
            if day_offset > learn_window:
                day_offset = learn_window

        add_task(t['topic'], t.get('chapter', ''), 'learn', t['difficulty'], hours, day_offset, t['priority_score'])
        day_load += hours

        revision_hours = max(0.5, round(hours * 0.4, 2))
        for r_type, r_gap in (('revision_1', 1), ('revision_2', 3), ('revision_3', 7)):
            r_offset = day_offset + r_gap
            if r_offset <= days_remaining - final_revision_days:
                add_task(t['topic'], t.get('chapter', ''), r_type, t['difficulty'], revision_hours, r_offset, t['priority_score'])

    # Weekly practice tests
    week = 1
    offset = 7
    while offset <= days_remaining - final_revision_days:
        add_task(f"Practice Test - Week {week}", '', 'practice_test', 'medium', 1.5, offset, 40)
        week += 1
        offset += 7

    # Final revision: spread all topics across the last `final_revision_days` days
    final_topics = [t['topic'] for t in ranked]
    if final_topics:
        chunks = [final_topics[i::final_revision_days] for i in range(final_revision_days)]
        for i, chunk in enumerate(chunks):
            f_offset = days_remaining - final_revision_days + i + 1
            for topic_name in chunk:
                add_task(topic_name, '', 'final_revision', 'medium', 0.5, f_offset, 90)

    meta = {
        'total_topics': len(topics),
        'total_hours': round(sum(t['estimated_hours'] for t in tasks), 1),
        'days_remaining': days_remaining,
        'warnings': warnings,
    }
    return tasks, meta


def generate_plan_summary(client, model, subject, exam_date, plan_type, meta):
    """Generate a short AI summary + tips for a newly created study plan."""
    prompt = f"""A student is preparing for a "{subject}" exam on {exam_date} using a {plan_type} study plan.
Total topics: {meta['total_topics']}, total estimated study hours: {meta['total_hours']}, days remaining: {meta['days_remaining']}.

Return a JSON object with:
- "summary": a 2-3 sentence encouraging overview of the plan
- "tips": array of 5 specific preparation tips for this subject

Return ONLY the JSON object, no other text."""
    system = "You are an expert academic planner. Respond with ONLY a valid JSON object. Do not include any explanation or thinking."
    try:
        return _ai_json_call(client, model, system, prompt)
    except Exception as e:
        logger.warning(f"Plan summary generation failed: {e}")
        return {'summary': f'A {plan_type} study plan for {subject} covering {meta["total_topics"]} topics.', 'tips': []}


def create_plan_shell(subject, exam_date, daily_hours, plan_type):
    conn = _get_conn()
    c = conn.cursor()
    c.execute('''INSERT INTO study_plans (subject, exam_date, hours_per_day, daily_hours, plan_type, plan_data)
        VALUES (?, ?, ?, ?, ?, ?)''', (subject, exam_date, daily_hours, daily_hours, plan_type, None))
    plan_id = c.lastrowid
    conn.commit()
    conn.close()
    return plan_id


def finalize_plan(plan_id, plan_data):
    conn = _get_conn()
    conn.execute("UPDATE study_plans SET plan_data = ? WHERE id = ?", (json.dumps(plan_data), plan_id))
    conn.commit()
    conn.close()


def save_tasks(plan_id, tasks):
    conn = _get_conn()
    c = conn.cursor()
    for t in tasks:
        c.execute('''INSERT INTO study_tasks
            (plan_id, topic_name, chapter_name, task_type, difficulty, estimated_hours, scheduled_date, priority_score)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)''',
            (plan_id, t['topic_name'], t.get('chapter_name', ''), t['task_type'], t['difficulty'],
             t['estimated_hours'], t['scheduled_date'], t.get('priority_score', 0)))
    conn.commit()
    conn.close()


def save_study_material(subject, plan_id, filename, original_name, file_type, pages, text, detected_topics):
    conn = _get_conn()
    c = conn.cursor()
    c.execute('''INSERT INTO uploaded_study_materials
        (plan_id, subject, file_name, original_name, file_type, pages, extracted_text, detected_topics)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)''',
        (plan_id, subject, filename, original_name, file_type, pages, text, json.dumps(detected_topics)))
    mat_id = c.lastrowid
    conn.commit()
    conn.close()
    return mat_id


def get_study_materials(plan_id=None, subject=None):
    conn = _get_conn()
    c = conn.cursor()
    query = "SELECT id, plan_id, subject, file_name, original_name, file_type, pages, detected_topics, upload_date FROM uploaded_study_materials WHERE 1=1"
    params = []
    if plan_id:
        query += " AND plan_id = ?"
        params.append(plan_id)
    if subject:
        query += " AND subject = ?"
        params.append(subject)
    query += " ORDER BY upload_date DESC"
    c.execute(query, params)
    rows = []
    for r in c.fetchall():
        d = dict(r)
        d['detected_topics'] = json.loads(d['detected_topics']) if d['detected_topics'] else []
        rows.append(d)
    conn.close()
    return rows


def get_study_plans():
    conn = _get_conn()
    c = conn.cursor()
    c.execute("SELECT * FROM study_plans ORDER BY created_at DESC")
    rows = []
    for r in c.fetchall():
        d = dict(r)
        d['plan_data'] = json.loads(d['plan_data']) if d['plan_data'] else {}
        rows.append(d)
    conn.close()
    return rows


def get_plan_tasks(plan_id):
    conn = _get_conn()
    c = conn.cursor()
    c.execute("SELECT * FROM study_tasks WHERE plan_id = ? ORDER BY scheduled_date ASC, priority_score DESC", (plan_id,))
    rows = [dict(r) for r in c.fetchall()]
    conn.close()
    return rows


def get_study_plan(plan_id):
    conn = _get_conn()
    c = conn.cursor()
    c.execute("SELECT * FROM study_plans WHERE id = ?", (plan_id,))
    row = c.fetchone()
    conn.close()
    if not row:
        return None
    d = dict(row)
    d['plan_data'] = json.loads(d['plan_data']) if d['plan_data'] else {}
    d['tasks'] = get_plan_tasks(plan_id)
    return d


def update_task_status(task_id, status):
    if status not in ('not_started', 'in_progress', 'completed'):
        status = 'not_started'
    conn = _get_conn()
    c = conn.cursor()
    c.execute("SELECT * FROM study_tasks WHERE id = ?", (task_id,))
    task = c.fetchone()
    if not task:
        conn.close()
        return None
    task = dict(task)
    c.execute("UPDATE study_tasks SET status = ?, updated_at = ? WHERE id = ?",
              (status, datetime.now().isoformat(), task_id))
    conn.commit()

    # Recompute topic completion percentage for this plan
    c.execute("SELECT status FROM study_tasks WHERE plan_id = ? AND topic_name = ?",
              (task['plan_id'], task['topic_name']))
    statuses = [r[0] for r in c.fetchall()]
    pct = round(100 * statuses.count('completed') / len(statuses), 1) if statuses else 0

    plan = get_study_plan(task['plan_id'])
    subject = plan['subject'] if plan else ''
    c.execute("SELECT id FROM study_progress WHERE plan_id = ? AND topic_name = ?", (task['plan_id'], task['topic_name']))
    existing = c.fetchone()
    if existing:
        c.execute("UPDATE study_progress SET completion_percentage = ?, last_updated = ? WHERE id = ?",
                  (pct, datetime.now().isoformat(), existing[0]))
    else:
        c.execute('''INSERT INTO study_progress (plan_id, topic_name, subject, completion_percentage, last_updated)
            VALUES (?, ?, ?, ?, ?)''', (task['plan_id'], task['topic_name'], subject, pct, datetime.now().isoformat()))
    conn.commit()
    conn.close()

    if status == 'completed' and plan:
        log_study_session(subject, task['topic_name'], int(task['estimated_hours'] * 60), task['task_type'])
    return pct


def rebalance_plan(plan_id):
    """Reschedule overdue, unfinished tasks into the earliest future days with capacity,
    without pushing anything past the exam date."""
    plan = get_study_plan(plan_id)
    if not plan:
        return 0

    today = datetime.now().date()
    try:
        exam_d = datetime.strptime(plan['exam_date'], '%Y-%m-%d').date()
    except ValueError:
        return 0

    tasks = plan['tasks']
    overdue = [t for t in tasks if t['status'] != 'completed' and t['scheduled_date']
               and datetime.strptime(t['scheduled_date'], '%Y-%m-%d').date() < today]
    if not overdue:
        return 0

    daily_budget = (plan.get('daily_hours') or plan.get('hours_per_day') or 2) * 1.2
    future_load = {}
    for t in tasks:
        if t['status'] == 'completed' or t in overdue:
            continue
        d = t['scheduled_date']
        future_load[d] = future_load.get(d, 0) + (t['estimated_hours'] or 0)

    overdue.sort(key=lambda x: x.get('priority_score', 0), reverse=True)
    conn = _get_conn()
    c = conn.cursor()
    moved = 0
    for t in overdue:
        day = today + timedelta(days=1)
        placed = False
        while day < exam_d:
            key = day.isoformat()
            load = future_load.get(key, 0)
            if load + t['estimated_hours'] <= daily_budget:
                future_load[key] = load + t['estimated_hours']
                c.execute("UPDATE study_tasks SET scheduled_date = ?, updated_at = ? WHERE id = ?",
                          (key, datetime.now().isoformat(), t['id']))
                moved += 1
                placed = True
                break
            day += timedelta(days=1)
        if not placed:
            # Best effort: place on the day before the exam anyway
            key = (exam_d - timedelta(days=1)).isoformat()
            future_load[key] = future_load.get(key, 0) + t['estimated_hours']
            c.execute("UPDATE study_tasks SET scheduled_date = ?, updated_at = ? WHERE id = ?",
                      (key, datetime.now().isoformat(), t['id']))
            moved += 1
    conn.commit()
    conn.close()
    return moved


def compute_readiness(plan_id):
    """Composite 0-100 readiness score from task completion, topic mastery, and overdue ratio."""
    plan = get_study_plan(plan_id)
    if not plan:
        return 0
    tasks = plan['tasks']
    today = datetime.now().date()
    total = len(tasks)
    if total == 0:
        return 0
    completed = sum(1 for t in tasks if t['status'] == 'completed')
    overdue = sum(1 for t in tasks if t['status'] != 'completed' and t['scheduled_date']
                  and datetime.strptime(t['scheduled_date'], '%Y-%m-%d').date() < today)

    completion_ratio = completed / total
    overdue_ratio = overdue / total

    conn = _get_conn()
    c = conn.cursor()
    c.execute("SELECT AVG(mastery_level) FROM topic_mastery WHERE subject = ?", (plan['subject'],))
    row = c.fetchone()
    conn.close()
    avg_mastery = (row[0] or 0) / 100 if row and row[0] is not None else completion_ratio

    readiness = 100 * (0.5 * completion_ratio + 0.3 * avg_mastery + 0.2 * (1 - overdue_ratio))
    return round(max(0, min(100, readiness)), 1)


def get_plan_summary_stats(plan_id):
    plan = get_study_plan(plan_id)
    if not plan:
        return None
    tasks = plan['tasks']
    today = datetime.now().date()
    total_topics = len({t['topic_name'] for t in tasks if t['task_type'] == 'learn'})
    total_hours = round(sum(t['estimated_hours'] for t in tasks), 1)
    completed = sum(1 for t in tasks if t['status'] == 'completed')
    total = len(tasks)
    completion_pct = round(100 * completed / total, 1) if total else 0
    try:
        exam_d = datetime.strptime(plan['exam_date'], '%Y-%m-%d').date()
        days_remaining = max(0, (exam_d - today).days)
    except ValueError:
        days_remaining = 0
    today_tasks = [t for t in tasks if t['scheduled_date'] == today.isoformat()]
    upcoming_tasks = [t for t in tasks if t['scheduled_date'] and today.isoformat() < t['scheduled_date'] <= (today + timedelta(days=3)).isoformat()]
    missed_tasks = [t for t in tasks if t['status'] != 'completed' and t['scheduled_date'] and t['scheduled_date'] < today.isoformat()]
    return {
        'total_topics': total_topics,
        'total_hours': total_hours,
        'days_remaining': days_remaining,
        'completion_pct': completion_pct,
        'today_tasks': today_tasks,
        'upcoming_tasks': upcoming_tasks,
        'missed_tasks': missed_tasks,
        'readiness': compute_readiness(plan_id),
    }


def group_tasks_for_display(tasks, plan_type):
    """Group learn tasks by day, week, or month for display, depending on plan_type."""
    groups = {}
    order = []
    for t in tasks:
        if not t['scheduled_date']:
            continue
        d = datetime.strptime(t['scheduled_date'], '%Y-%m-%d').date()
        if plan_type == 'weekly':
            week_start = d - timedelta(days=d.weekday())
            key = f"Week of {week_start.strftime('%b %d, %Y')}"
        elif plan_type == 'monthly':
            key = d.strftime('%B %Y')
        else:
            key = d.strftime('%a, %b %d, %Y')
        if key not in groups:
            groups[key] = []
            order.append(key)
        groups[key].append(t)
    return [(key, groups[key]) for key in order]


def get_dashboard_study_overview():
    """Today/upcoming/missed tasks + readiness across all active plans, for the dashboard widget."""
    conn = _get_conn()
    c = conn.cursor()
    c.execute("SELECT id FROM study_plans WHERE status = 'active' ORDER BY created_at DESC")
    plan_ids = [r[0] for r in c.fetchall()]
    conn.close()

    today_tasks, upcoming_tasks, missed_tasks, readiness_scores = [], [], [], []
    for pid in plan_ids:
        stats = get_plan_summary_stats(pid)
        if not stats:
            continue
        today_tasks.extend(stats['today_tasks'])
        upcoming_tasks.extend(stats['upcoming_tasks'])
        missed_tasks.extend(stats['missed_tasks'])
        readiness_scores.append(stats['readiness'])

    avg_readiness = round(sum(readiness_scores) / len(readiness_scores), 1) if readiness_scores else 0
    return {
        'today_tasks': today_tasks[:8],
        'upcoming_tasks': upcoming_tasks[:8],
        'missed_tasks': missed_tasks[:8],
        'readiness': avg_readiness,
    }


def create_smart_study_plan(client, model, subject, exam_date, daily_hours, plan_type, materials, manual_topics_text):
    """Orchestrates the full smart planning workflow: detect/classify topics from uploaded
    materials (or manual input), prioritize using exam predictions + weak topics, build the
    rule-based schedule, and persist everything. Returns the new plan_id.

    materials: list of dicts {filename, original_name, file_type, pages, text}
    """
    plan_id = create_plan_shell(subject, exam_date, daily_hours, plan_type)

    candidate_topics = []
    context_snippet = None
    if materials:
        for m in materials:
            file_topics = extract_material_topics(client, model, subject, m['text'])
            save_study_material(subject, plan_id, m['filename'], m['original_name'],
                               m['file_type'], m['pages'], m['text'], file_topics)
            candidate_topics.extend(file_topics)
        context_snippet = materials[0]['text']
    else:
        for part in re.split(r'[,\n]', manual_topics_text or ''):
            part = part.strip()
            if part:
                candidate_topics.append(part)

    candidate_topics = list(dict.fromkeys(candidate_topics))
    if not candidate_topics:
        finalize_plan(plan_id, {
            'summary': '', 'tips': [], 'total_topics': 0, 'total_hours': 0, 'days_remaining': 0,
            'warnings': ['No topics could be determined from the uploaded materials or manual input.']
        })
        return plan_id

    classified = classify_topics(client, model, subject, candidate_topics, context_snippet)

    boost_map = _prediction_priority_boost(subject)
    weak = [w['topic'] for w in detect_weak_topics() if w.get('subject') == subject]
    for t in classified:
        t['boost'] = _match_boost(t['topic'], boost_map)

    tasks, meta = build_study_schedule(classified, exam_date, daily_hours, plan_type, weak)
    save_tasks(plan_id, tasks)

    summary_tips = generate_plan_summary(client, model, subject, exam_date, plan_type, meta)
    finalize_plan(plan_id, {**meta, **summary_tips})
    return plan_id


# ─── Progress Tracking ───────────────────────────────────────────────

def log_study_session(subject, topic, duration_minutes, activity_type):
    conn = _get_conn()
    c = conn.cursor()
    c.execute('''INSERT INTO study_sessions (subject, topic, duration_minutes, activity_type)
        VALUES (?, ?, ?, ?)''', (subject, topic, duration_minutes, activity_type))
    conn.commit()
    conn.close()
    if subject and topic:
        _update_mastery(subject, topic, None, studied=True)


def get_progress_data():
    conn = _get_conn()
    c = conn.cursor()

    # Total study time
    c.execute("SELECT COALESCE(SUM(duration_minutes), 0) FROM study_sessions")
    total_minutes = c.fetchone()[0]

    # Total quizzes & avg score
    c.execute("SELECT COUNT(*), COALESCE(AVG(percentage), 0) FROM quiz_attempts")
    row = c.fetchone()
    total_quizzes = row[0]
    avg_score = round(row[1], 1)

    # Flashcard stats
    c.execute("SELECT COUNT(*) FROM flashcards")
    total_flashcards = c.fetchone()[0]
    c.execute("SELECT COUNT(*) FROM flashcards WHERE confidence >= 4")
    mastered_cards = c.fetchone()[0]

    # Topics mastered vs weak
    c.execute("SELECT COUNT(*) FROM topic_mastery WHERE mastery_level >= 70")
    strong_topics = c.fetchone()[0]
    c.execute("SELECT COUNT(*) FROM topic_mastery WHERE is_weak = 1")
    weak_topics = c.fetchone()[0]

    # Study sessions by activity
    c.execute("SELECT activity_type, COUNT(*), SUM(duration_minutes) FROM study_sessions GROUP BY activity_type")
    activities = [dict(r) for r in c.fetchall()]

    # Daily study (last 7 days)
    c.execute("""SELECT DATE(started_at) as day, SUM(duration_minutes) as mins
        FROM study_sessions GROUP BY day ORDER BY day DESC LIMIT 7""")
    daily_study = [dict(r) for r in c.fetchall()]

    # Recent quiz scores
    c.execute("""SELECT qa.percentage, qa.attempted_at, q.topic 
        FROM quiz_attempts qa JOIN quizzes q ON qa.quiz_id = q.id
        ORDER BY qa.attempted_at DESC LIMIT 10""")
    recent_scores = [dict(r) for r in c.fetchall()]

    # Subject breakdown
    c.execute("""SELECT subject, COUNT(*) as sessions, SUM(duration_minutes) as mins
        FROM study_sessions WHERE subject IS NOT NULL
        GROUP BY subject ORDER BY mins DESC""")
    subject_breakdown = [dict(r) for r in c.fetchall()]

    conn.close()

    return {
        'total_study_minutes': total_minutes,
        'total_study_hours': round(total_minutes / 60, 1),
        'total_quizzes': total_quizzes,
        'avg_quiz_score': avg_score,
        'total_flashcards': total_flashcards,
        'mastered_cards': mastered_cards,
        'strong_topics': strong_topics,
        'weak_topics': weak_topics,
        'activities': activities,
        'daily_study': daily_study,
        'recent_scores': recent_scores,
        'subject_breakdown': subject_breakdown
    }


def _update_mastery(subject, topic, score=None, studied=False):
    conn = _get_conn()
    c = conn.cursor()
    c.execute("SELECT * FROM topic_mastery WHERE subject = ? AND topic = ?", (subject, topic))
    row = c.fetchone()

    if row:
        row = dict(row)
        scores = json.loads(row['quiz_scores']) if row['quiz_scores'] else []
        times = row['times_studied']
        if score is not None:
            scores.append(score)
        if studied:
            times += 1
        avg = sum(scores) / len(scores) if scores else 0
        is_weak = 1 if (len(scores) >= 2 and avg < 50) else 0
        c.execute("""UPDATE topic_mastery SET 
            mastery_level = ?, quiz_scores = ?, times_studied = ?,
            is_weak = ?, last_studied = ?, updated_at = ?
            WHERE id = ?""",
            (round(avg, 1), json.dumps(scores), times, is_weak,
             datetime.now().isoformat(), datetime.now().isoformat(), row['id']))
    else:
        scores = [score] if score is not None else []
        avg = score if score is not None else 0
        is_weak = 1 if (score is not None and score < 50) else 0
        c.execute("""INSERT INTO topic_mastery 
            (subject, topic, mastery_level, quiz_scores, times_studied, is_weak, last_studied, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
            (subject, topic, avg, json.dumps(scores), 1 if studied else 0,
             is_weak, datetime.now().isoformat(), datetime.now().isoformat()))

    conn.commit()
    conn.close()


def detect_weak_topics():
    conn = _get_conn()
    c = conn.cursor()
    c.execute("""SELECT subject, topic, mastery_level, quiz_scores, times_studied
        FROM topic_mastery WHERE is_weak = 1 ORDER BY mastery_level ASC""")
    weak = [dict(r) for r in c.fetchall()]
    for w in weak:
        w['quiz_scores'] = json.loads(w['quiz_scores']) if w['quiz_scores'] else []
    conn.close()
    return weak


def get_all_mastery():
    conn = _get_conn()
    c = conn.cursor()
    c.execute("SELECT * FROM topic_mastery ORDER BY mastery_level ASC")
    rows = []
    for r in c.fetchall():
        d = dict(r)
        d['quiz_scores'] = json.loads(d['quiz_scores']) if d['quiz_scores'] else []
        rows.append(d)
    conn.close()
    return rows


# ─── Notes ────────────────────────────────────────────────────────────

def convert_to_notes(client, model, lecture_text, subject, title):
    """Convert lecture text to structured notes using AI."""
    prompt = f"""Convert the following lecture content into well-structured study notes for the subject "{subject}".

Title: {title}

Requirements:
- Start with a brief summary (3-4 sentences)
- Use clear headers and subheaders  
- Use bullet points for key concepts
- Highlight important terms with **bold**
- Include a "Key Takeaways" section at the end
- Add a "Quick Review Questions" section with 3-5 questions

Lecture Content:
{lecture_text[:5000]}

Return the notes in clean markdown format."""

    raw = _ai_call(client, model,
        "You are an expert note-taker. Create clear, comprehensive study notes from lectures. Do not use thinking tags.",
        prompt)
    # Clean any residual thinking tags
    raw = re.sub(r'<think>[\s\S]*?</think>', '', raw).strip()
    return raw


def save_notes(subject, title, content, source_type='lecture'):
    conn = _get_conn()
    c = conn.cursor()
    c.execute('''INSERT INTO notes (subject, title, content, source_type)
        VALUES (?, ?, ?, ?)''', (subject, title, content, source_type))
    note_id = c.lastrowid
    conn.commit()
    conn.close()
    return note_id


def get_all_notes():
    conn = _get_conn()
    c = conn.cursor()
    c.execute("SELECT id, subject, title, source_type, created_at FROM notes ORDER BY created_at DESC")
    rows = [dict(r) for r in c.fetchall()]
    conn.close()
    return rows


def get_note(note_id):
    conn = _get_conn()
    c = conn.cursor()
    c.execute("SELECT * FROM notes WHERE id = ?", (note_id,))
    row = c.fetchone()
    conn.close()
    return dict(row) if row else None


# ─── Exam Prediction Assistant ───────────────────────────────────────

ANALYSIS_TYPES = ('frequent_topics', 'repeated_questions', 'high_weightage', 'full')


def save_exam_document(subject, doc_type, filename, original_name, text, year=None):
    """Save an uploaded exam document (paper/question bank) and index it into Chroma."""
    import rag
    conn = _get_conn()
    c = conn.cursor()
    c.execute('''INSERT INTO exam_documents (subject, doc_type, filename, original_name, year, extracted_text)
        VALUES (?, ?, ?, ?, ?, ?)''', (subject, doc_type, filename, original_name, year, text))
    doc_id = c.lastrowid
    conn.commit()

    # Embedding/indexing happens outside any open transaction on qa_data.db
    chunk_count = rag.index_document(doc_id, subject, doc_type, text, year)
    if chunk_count:
        c.execute("UPDATE exam_documents SET chunk_count = ? WHERE id = ?", (chunk_count, doc_id))
        conn.commit()
    conn.close()
    return doc_id


def get_exam_documents(subject=None, doc_type=None):
    conn = _get_conn()
    c = conn.cursor()
    query = "SELECT id, subject, doc_type, filename, original_name, year, chunk_count, uploaded_at FROM exam_documents WHERE 1=1"
    params = []
    if subject:
        query += " AND subject = ?"
        params.append(subject)
    if doc_type:
        query += " AND doc_type = ?"
        params.append(doc_type)
    query += " ORDER BY uploaded_at DESC"
    c.execute(query, params)
    rows = [dict(r) for r in c.fetchall()]
    conn.close()
    return rows


def _build_rag_context(subject, syllabus_topics=None, max_chars=9000):
    """Retrieve the most relevant chunks from Chroma for a subject via a few varied queries."""
    import rag
    queries = [
        f"important exam questions and topics in {subject}",
        f"frequently repeated questions in {subject} exams",
        f"key concepts and definitions in {subject}",
    ]
    if syllabus_topics:
        queries.append(syllabus_topics[:500])

    seen = set()
    pieces = []
    total = 0
    for q in queries:
        for chunk in rag.query_context(subject, q, top_k=6):
            if chunk in seen:
                continue
            seen.add(chunk)
            pieces.append(chunk)
            total += len(chunk)
            if total >= max_chars:
                break
        if total >= max_chars:
            break
    return "\n\n---\n\n".join(pieces)[:max_chars]


def _ai_json_call(client, model, system, prompt):
    """Call the AI expecting a JSON object/array back, with retries on parse failure."""
    for attempt in range(3):
        try:
            raw = _ai_call(client, model, system, prompt)
            return _parse_json_response(raw)
        except (json.JSONDecodeError, ValueError) as e:
            logger.warning(f"AI JSON call parse attempt {attempt + 1} failed: {e}")
            if attempt == 2:
                raise
            time.sleep(0.5)


def analyze_exam_documents(client, model, subject, analysis_type='full', syllabus_topics=None):
    """Analyze previously uploaded exam documents for a subject using RAG + AI to
    predict topics, repeated questions, weightage, predicted questions/paper, and tips.

    Split into two smaller AI calls (analysis, then predicted questions/paper) since a
    single call asking for everything at once tends to produce truncated/invalid JSON.
    """
    if analysis_type not in ANALYSIS_TYPES:
        analysis_type = 'full'

    context = _build_rag_context(subject, syllabus_topics)
    if not context.strip():
        raise ValueError(
            "No exam documents found for this subject. Please upload previous year "
            "papers or question banks before running the analysis.")

    syllabus_line = f"\nSyllabus topics provided by the student: {syllabus_topics}" if syllabus_topics else ""
    focus = {
        'frequent_topics': "Focus mainly on frequently asked topics, but still fill in every field below as best you can.",
        'repeated_questions': "Focus mainly on repeated questions, but still fill in every field below as best you can.",
        'high_weightage': "Focus mainly on high weightage concepts, but still fill in every field below as best you can.",
        'full': "Provide a thorough analysis across every field below.",
    }[analysis_type]

    system = ("You are an expert exam analyst who studies historical exam papers to identify patterns and predict "
              "future exam content. Respond with ONLY valid JSON. Keep answers concise. Do not include any "
              "explanation, commentary, or thinking.")

    # Call 1: topic/question pattern analysis
    analysis_prompt = f"""You are analyzing previous year exam papers and question banks for the subject "{subject}" to identify exam patterns.{syllabus_line}

Relevant excerpts from the uploaded documents:
{context}

{focus}

Return a JSON object with exactly these fields:
- "frequently_asked_topics": array of up to 8 objects, each with "topic", "frequency" (number of times referenced), "importance_score" (0-100), "confidence" (0-100)
- "repeated_questions": array of up to 8 objects, each with "question", "years_appeared" (array of years if identifiable, else []), "importance" (one of "Highly Important", "Important", "Moderate")
- "high_weightage_concepts": array of up to 8 objects, each with "concept", "exam_frequency" (0-100), "weightage_level" (one of "Critical", "High", "Medium", "Low")
- "important_chapters": array of up to 8 chapter/unit names ranked by importance
- "preparation_tips": object with "priority_topics" (array of up to 5), "chapters_to_revise_first" (array of up to 5), "important_formulas" (array of up to 5), "frequently_tested_concepts" (array of up to 5)

Keep each string concise (under 25 words). Return ONLY the JSON object, no other text."""

    result = dict(_ai_json_call(client, model, system, analysis_prompt))

    # Call 2: predicted questions by mark value (flat schema - easier for the model to follow)
    top_topics = ', '.join(t.get('topic', '') for t in result.get('frequently_asked_topics', [])[:6]) or subject
    questions_prompt = f"""Based on these frequently examined topics in "{subject}": {top_topics}.{syllabus_line}

Relevant excerpts from the uploaded documents:
{context[:4000]}

Generate exam-style predicted questions grouped by mark value.

Return a JSON object with exactly these fields (each an array of question strings only, no answers):
- "1_mark": 5 very short-answer/definition questions
- "2_mark": 5 short-explanation questions
- "5_mark": 5 detailed/descriptive questions
- "10_mark": 3 long analytical questions

Keep each question concise. Return ONLY the JSON object, no other text."""

    result['predicted_questions'] = _ai_json_call(client, model, system, questions_prompt)

    # Call 3: a model predicted exam paper (flat schema)
    paper_prompt = f"""Based on these frequently examined topics in "{subject}": {top_topics}.{syllabus_line}

Relevant excerpts from the uploaded documents:
{context[:4000]}

Generate a model predicted exam paper with three parts.

Return a JSON object with exactly these fields (each an array of objects with "question" and "answer"):
- "part_a": 3 short-answer questions with brief answers
- "part_b": 3 medium-length questions with concise model answers
- "part_c": 2 long-answer questions with detailed model answers

Keep answers concise. Return ONLY the JSON object, no other text."""

    result['predicted_paper'] = _ai_json_call(client, model, system, paper_prompt)
    return result


def save_full_prediction(subject, analysis_type, data):
    """Persist a full prediction result: the raw JSON blob plus normalized rows
    in topic_predictions / repeated_questions / predicted_papers for querying."""
    conn = _get_conn()
    c = conn.cursor()
    c.execute('''INSERT INTO predictions (subject, prediction_data, analysis_type)
        VALUES (?, ?, ?)''', (subject, json.dumps(data), analysis_type))
    pred_id = c.lastrowid

    for t in data.get('frequently_asked_topics', []) or []:
        c.execute('''INSERT INTO topic_predictions (prediction_id, subject, topic, frequency_score, confidence_score, importance_level)
            VALUES (?, ?, ?, ?, ?, ?)''',
            (pred_id, subject, t.get('topic', ''), t.get('importance_score', 0),
             t.get('confidence', 0), t.get('importance', 'medium')))

    for rq in data.get('repeated_questions', []) or []:
        c.execute('''INSERT INTO repeated_questions (prediction_id, subject, question, years_found, importance)
            VALUES (?, ?, ?, ?, ?)''',
            (pred_id, subject, rq.get('question', ''),
             json.dumps(rq.get('years_appeared', [])), rq.get('importance', 'medium')))

    if data.get('predicted_paper'):
        c.execute('''INSERT INTO predicted_papers (prediction_id, subject, paper_data)
            VALUES (?, ?, ?)''', (pred_id, subject, json.dumps(data['predicted_paper'])))

    conn.commit()
    conn.close()
    return pred_id


def get_predictions(subject=None):
    conn = _get_conn()
    c = conn.cursor()
    if subject:
        c.execute("SELECT * FROM predictions WHERE subject = ? ORDER BY created_at DESC LIMIT 10", (subject,))
    else:
        c.execute("SELECT * FROM predictions ORDER BY created_at DESC LIMIT 10")
    rows = []
    for r in c.fetchall():
        d = dict(r)
        d['prediction_data'] = json.loads(d['prediction_data']) if d['prediction_data'] else {}
        rows.append(d)
    conn.close()
    return rows


def get_prediction(prediction_id):
    conn = _get_conn()
    c = conn.cursor()
    c.execute("SELECT * FROM predictions WHERE id = ?", (prediction_id,))
    row = c.fetchone()
    conn.close()
    if not row:
        return None
    d = dict(row)
    d['prediction_data'] = json.loads(d['prediction_data']) if d['prediction_data'] else {}
    return d


# ─── Practice Test Generator ─────────────────────────────────────────

PRACTICE_TEST_TYPES = {'mcq': 'mcq', 'short': '2mark', 'long': '5mark'}


def generate_practice_test(client, model, prediction, test_type='mcq', count=10, difficulty='medium'):
    """Build a practice test (MCQ / short-answer / long-answer) from a prediction's
    top topics, reusing the quiz generation/storage pipeline."""
    question_type = PRACTICE_TEST_TYPES.get(test_type, 'mcq')
    data = prediction.get('prediction_data', {})
    subject = prediction.get('subject', 'General')

    topics = [t.get('topic', '') for t in data.get('frequently_asked_topics', []) if t.get('topic')]
    chapters = data.get('important_chapters', [])
    topic_line = ', '.join(topics[:8]) or subject
    context_parts = []
    if chapters:
        context_parts.append("Important chapters: " + ', '.join(chapters))
    if data.get('high_weightage_concepts'):
        concepts = [c.get('concept', '') for c in data['high_weightage_concepts'] if c.get('concept')]
        if concepts:
            context_parts.append("High weightage concepts: " + ', '.join(concepts))
    context = '\n'.join(context_parts) or None

    questions = generate_exam_quiz(client, model, topic_line, question_type, difficulty, count, context)
    quiz_id = save_quiz(subject, f"Practice Test: {topic_line[:60]}", difficulty, questions, question_type)
    return quiz_id


# ─── Export (Predictions & Practice Tests) ───────────────────────────

def _escape_pdf_text(text):
    text = str(text).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    return text.replace('\n', '<br/>')


def _write_pdf_report(title, sections, filepath):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem

    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(filepath, pagesize=A4)
    flow = [Paragraph(_escape_pdf_text(title), styles['Title']), Spacer(1, 16)]
    for heading, lines in sections:
        flow.append(Paragraph(_escape_pdf_text(heading), styles['Heading2']))
        flow.append(Spacer(1, 6))
        items = [ListItem(Paragraph(_escape_pdf_text(line), styles['Normal'])) for line in lines if str(line).strip()]
        if items:
            flow.append(ListFlowable(items, bulletType='bullet'))
        flow.append(Spacer(1, 12))
    doc.build(flow)


def _write_txt_report(title, sections, filepath):
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(title + '\n' + '=' * len(title) + '\n\n')
        for heading, lines in sections:
            f.write(heading + '\n' + '-' * len(heading) + '\n')
            for line in lines:
                if str(line).strip():
                    f.write(f"- {line}\n")
            f.write('\n')


def _prediction_to_sections(prediction):
    data = prediction.get('prediction_data', {})
    sections = []

    if data.get('frequently_asked_topics'):
        lines = [f"{t.get('topic')} — frequency {t.get('frequency')}, importance score {t.get('importance_score')}, confidence {t.get('confidence')}%"
                 for t in data['frequently_asked_topics']]
        sections.append(('Frequently Asked Topics', lines))

    if data.get('repeated_questions'):
        lines = [f"{q.get('question')} (years: {', '.join(str(y) for y in q.get('years_appeared', []))}) — {q.get('importance')}"
                 for q in data['repeated_questions']]
        sections.append(('Repeated Questions', lines))

    if data.get('high_weightage_concepts'):
        lines = [f"{c.get('concept')} — exam frequency {c.get('exam_frequency')}%, weightage {c.get('weightage_level')}"
                 for c in data['high_weightage_concepts']]
        sections.append(('High Weightage Concepts', lines))

    if data.get('important_chapters'):
        sections.append(('Important Chapters', data['important_chapters']))

    if data.get('predicted_questions'):
        pq = data['predicted_questions']
        for key, label in (('1_mark', '1 Mark Questions'), ('2_mark', '2 Mark Questions'),
                            ('5_mark', '5 Mark Questions'), ('10_mark', '10 Mark Questions')):
            if pq.get(key):
                sections.append((label, pq[key]))

    if data.get('predicted_paper'):
        pp = data['predicted_paper']
        for key, label in (('part_a', 'Predicted Paper — Part A'), ('part_b', 'Predicted Paper — Part B'),
                            ('part_c', 'Predicted Paper — Part C')):
            if pp.get(key):
                lines = [f"Q: {item.get('question')}\nA: {item.get('answer')}" for item in pp[key]]
                sections.append((label, lines))

    if data.get('preparation_tips'):
        tips = data['preparation_tips']
        for key, label in (('priority_topics', 'Priority Topics'), ('chapters_to_revise_first', 'Chapters to Revise First'),
                            ('important_formulas', 'Important Formulas'), ('frequently_tested_concepts', 'Frequently Tested Concepts')):
            if tips.get(key):
                sections.append((label, tips[key]))

    return sections


def export_prediction(prediction, fmt='txt'):
    """Export a prediction report to TXT or PDF. Returns the file path."""
    sections = _prediction_to_sections(prediction)
    title = f"Exam Prediction Report - {prediction.get('subject', '')}"
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    if fmt == 'pdf':
        filepath = os.path.join(EXPORT_DIR, f"prediction_{prediction['id']}_{timestamp}.pdf")
        _write_pdf_report(title, sections, filepath)
    else:
        filepath = os.path.join(EXPORT_DIR, f"prediction_{prediction['id']}_{timestamp}.txt")
        _write_txt_report(title, sections, filepath)
    return filepath


def export_quiz(quiz_id, fmt='txt'):
    """Export a quiz/practice test to TXT or PDF. Returns the file path."""
    quiz, questions = get_quiz(quiz_id)
    if not quiz:
        raise ValueError('Quiz not found')

    lines = []
    for i, q in enumerate(questions, 1):
        if q['question_type'] == 'mcq':
            lines.append(
                f"Q{i}. {q['question']}\n"
                f"A) {q['option_a']}   B) {q['option_b']}   C) {q['option_c']}   D) {q['option_d']}\n"
                f"Correct Answer: {q['correct_answer']}\n"
                f"Explanation: {q.get('explanation', '')}"
            )
        else:
            lines.append(
                f"Q{i}. {q['question']} [{q['marks']} mark{'s' if q['marks'] != 1 else ''}]\n"
                f"Model Answer: {q.get('model_answer', '')}"
            )

    sections = [(f"{quiz['topic']} ({quiz['question_type'].upper()})", lines)]
    title = f"Practice Test - {quiz.get('subject', '')} - {quiz['topic']}"
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    if fmt == 'pdf':
        filepath = os.path.join(EXPORT_DIR, f"quiz_{quiz_id}_{timestamp}.pdf")
        _write_pdf_report(title, sections, filepath)
    else:
        filepath = os.path.join(EXPORT_DIR, f"quiz_{quiz_id}_{timestamp}.txt")
        _write_txt_report(title, sections, filepath)
    return filepath


# ─── Dashboard ────────────────────────────────────────────────────────

def get_dashboard_data():
    conn = _get_conn()
    c = conn.cursor()

    # Quick stats
    c.execute("SELECT COUNT(*) FROM quizzes")
    total_quizzes = c.fetchone()[0]
    c.execute("SELECT COUNT(*) FROM flashcards")
    total_flashcards = c.fetchone()[0]
    c.execute("SELECT COUNT(*) FROM notes")
    total_notes = c.fetchone()[0]
    c.execute("SELECT COUNT(*) FROM pdfs")
    total_pdfs = c.fetchone()[0]
    c.execute("SELECT COALESCE(SUM(duration_minutes), 0) FROM study_sessions")
    total_study_mins = c.fetchone()[0]
    c.execute("SELECT COUNT(*) FROM topic_mastery WHERE is_weak = 1")
    weak_count = c.fetchone()[0]
    c.execute("SELECT COUNT(*) FROM study_plans WHERE status = 'active'")
    active_plans = c.fetchone()[0]

    # Recent activity
    c.execute("""SELECT 'quiz' as type, q.topic as title, qa.percentage as detail, qa.attempted_at as ts
        FROM quiz_attempts qa JOIN quizzes q ON qa.quiz_id = q.id
        UNION ALL
        SELECT 'flashcard' as type, topic as title, CAST(COUNT(*) AS TEXT) as detail, MAX(created_at) as ts
        FROM flashcards GROUP BY topic
        UNION ALL
        SELECT 'note' as type, title, source_type as detail, created_at as ts FROM notes
        ORDER BY ts DESC LIMIT 8""")
    recent_activity = [dict(r) for r in c.fetchall()]

    # Due flashcards
    now = datetime.now().isoformat()
    c.execute("SELECT COUNT(*) FROM flashcards WHERE next_review IS NULL OR next_review <= ?", (now,))
    due_flashcards = c.fetchone()[0]

    conn.close()

    return {
        'total_quizzes': total_quizzes,
        'total_flashcards': total_flashcards,
        'total_notes': total_notes,
        'total_pdfs': total_pdfs,
        'total_study_hours': round(total_study_mins / 60, 1),
        'weak_topics': weak_count,
        'active_plans': active_plans,
        'due_flashcards': due_flashcards,
        'recent_activity': recent_activity
    }
